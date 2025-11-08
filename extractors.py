"""
Text extraction layer - Extract text from different file formats
"""
import os
import logging
from typing import Optional
from pathlib import Path
import PyPDF2
from docx import Document

# Setup logging
logger = logging.getLogger(__name__)


class TextExtractor:
    """Extract text from various file formats"""
    
    @staticmethod
    def extract_from_pdf(file_path: str) -> str:
        """
        Extract text from PDF file with enhanced handling for encrypted and scanned PDFs
        
        Args:
            file_path: Path to the PDF file
            
        Returns:
            Extracted text content
            
        Raises:
            ValueError: If PDF is encrypted and cannot be decrypted, or extraction fails
        """
        try:
            logger.info(f"Extracting text from PDF: {file_path}")
            text_content = []
            
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                # Handle encrypted PDFs
                if pdf_reader.is_encrypted:
                    logger.warning(f"PDF is encrypted: {file_path}")
                    try:
                        # Try to decrypt with empty password (common case)
                        pdf_reader.decrypt("")
                        logger.info("Successfully decrypted PDF with empty password")
                    except Exception as decrypt_error:
                        raise ValueError(
                            "PDF is password-protected and cannot be decrypted. "
                            "Please provide an unencrypted version of the file."
                        )
                
                num_pages = len(pdf_reader.pages)
                logger.info(f"PDF has {num_pages} pages")
                
                has_text = False
                for page_num in range(num_pages):
                    page = pdf_reader.pages[page_num]
                    text = page.extract_text()
                    
                    if text and text.strip():
                        text_content.append(text.strip())
                        has_text = True
                
                # Handle scanned/image-only PDFs
                if not has_text:
                    logger.warning(f"No extractable text found in PDF: {file_path}")
                    text_content.append(
                        "[Warning: This appears to be a scanned or image-only PDF. "
                        "No text could be extracted. Please provide a text-based PDF or use OCR.]"
                    )
            
            # Clean and join text with consistent formatting
            clean_text = "\n".join(line.strip() for line in text_content if line.strip())
            
            logger.info(f"Successfully extracted {len(clean_text)} characters from PDF")
            return clean_text
            
        except ValueError:
            # Re-raise ValueError with our custom messages
            raise
        except Exception as e:
            logger.error(f"Error extracting text from PDF: {str(e)}")
            raise ValueError(f"Error extracting text from PDF: {str(e)}")
    
    @staticmethod
    def extract_from_docx(file_path: str) -> str:
        """
        Extract text from DOCX file with improved formatting
        
        Args:
            file_path: Path to the DOCX file
            
        Returns:
            Extracted text content
        """
        try:
            logger.info(f"Extracting text from DOCX: {file_path}")
            doc = Document(file_path)
            text_content = []
            
            # Extract paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text.strip())
            
            # Extract tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_content.append(cell.text.strip())
            
            # Clean and join text
            clean_text = "\n".join(line for line in text_content if line)
            
            logger.info(f"Successfully extracted {len(clean_text)} characters from DOCX")
            return clean_text
            
        except Exception as e:
            logger.error(f"Error extracting text from DOCX: {str(e)}")
            raise ValueError(f"Error extracting text from DOCX: {str(e)}")
    
    @staticmethod
    def extract_from_txt(file_path: str) -> str:
        """
        Extract text from TXT file with multiple encoding support
        
        Args:
            file_path: Path to the TXT file
            
        Returns:
            Extracted text content
        """
        try:
            logger.info(f"Extracting text from TXT: {file_path}")
            encodings = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']
            
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as file:
                        text = file.read()
                        logger.info(f"Successfully read TXT file with {encoding} encoding")
                        
                        # Clean up excessive whitespace
                        lines = [line.strip() for line in text.split('\n')]
                        clean_text = "\n".join(line for line in lines if line)
                        
                        return clean_text
                except UnicodeDecodeError:
                    continue
            
            raise ValueError(
                f"Unable to decode file with common encodings: {', '.join(encodings)}"
            )
        except ValueError:
            raise
        except Exception as e:
            logger.error(f"Error extracting text from TXT: {str(e)}")
            raise ValueError(f"Error extracting text from TXT: {str(e)}")
    
    @staticmethod
    def extract_from_doc(file_path: str) -> str:
        """
        Extract text from legacy DOC file (MS Word 97-2003)
        
        Args:
            file_path: Path to the DOC file
            
        Returns:
            Extracted text content
            
        Note:
            Requires textract library: pip install textract
            For Windows, you may also need antiword
        """
        try:
            logger.info(f"Extracting text from DOC: {file_path}")
            import textract
            
            text = textract.process(file_path).decode('utf-8')
            
            # Clean up text
            lines = [line.strip() for line in text.split('\n')]
            clean_text = "\n".join(line for line in lines if line)
            
            logger.info(f"Successfully extracted {len(clean_text)} characters from DOC")
            return clean_text
            
        except ImportError:
            logger.error("textract library not installed")
            raise ValueError(
                "Legacy .doc file format requires 'textract' library. "
                "Please install it: pip install textract"
            )
        except Exception as e:
            logger.error(f"Error extracting text from DOC: {str(e)}")
            raise ValueError(f"Error extracting text from DOC: {str(e)}")
    
    @classmethod
    def extract(cls, file_path: str) -> str:
        """
        Extract text from file based on extension
        
        Args:
            file_path: Path to the file
            
        Returns:
            Extracted text content
        """
        extension = Path(file_path).suffix.lower()
        
        logger.info(f"Starting text extraction for file: {file_path} (type: {extension})")
        
        if extension == '.pdf':
            return cls.extract_from_pdf(file_path)
        elif extension == '.docx':
            return cls.extract_from_docx(file_path)
        elif extension == '.doc':
            return cls.extract_from_doc(file_path)
        elif extension == '.txt':
            return cls.extract_from_txt(file_path)
        else:
            raise ValueError(f"Unsupported file format: {extension}")
    
    @staticmethod
    def extract_from_pdf_generator(file_path: str):
        """
        Generator-based PDF extraction for large files (streaming approach)
        
        Args:
            file_path: Path to the PDF file
            
        Yields:
            Text content from each page
            
        Usage:
            for page_text in TextExtractor.extract_from_pdf_generator('large.pdf'):
                process(page_text)
        """
        try:
            logger.info(f"Starting streaming extraction from PDF: {file_path}")
            
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                # Handle encryption
                if pdf_reader.is_encrypted:
                    try:
                        pdf_reader.decrypt("")
                    except:
                        raise ValueError("PDF is password-protected")
                
                num_pages = len(pdf_reader.pages)
                logger.info(f"Streaming {num_pages} pages")
                
                for page_num, page in enumerate(pdf_reader.pages):
                    text = page.extract_text()
                    if text and text.strip():
                        yield text.strip()
                    else:
                        yield f"[Page {page_num + 1}: No extractable text]"
                        
        except Exception as e:
            logger.error(f"Error in streaming PDF extraction: {str(e)}")
            raise ValueError(f"Error extracting PDF: {str(e)}")
